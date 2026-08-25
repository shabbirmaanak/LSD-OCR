"""
Non-AI Document Parser for .docx, .txt, and .pdf files.
Extracts raw text deterministically from user uploaded documents with 100% font/text fidelity.
Supports PyMuPDF high-resolution page rendering and offline OCR fallback for legacy custom font PDFs.
"""

import io
import re
from typing import List, Tuple, Dict, Any, Optional
from docx import Document
import pypdf

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Parses document bytes into plain text string.
    Supports .docx, .txt, .pdf, and scanned image formats.
    Auto-detects corrupted custom-font PDF streams and falls back to PyMuPDF page rendering + OCR.
    """
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    if ext == 'docx':
        return parse_docx(file_bytes)
    elif ext == 'pdf':
        return parse_pdf(file_bytes)
    elif ext in ['png', 'jpg', 'jpeg', 'webp', 'bmp', 'tiff']:
        try:
            from backend.ocr_engine import extract_text_from_image
            res = extract_text_from_image(file_bytes)
            if res.get("success"):
                return res.get("text", "")
        except Exception as e:
            print(f"Image OCR error: {e}")
        return ""
    else:
        return parse_txt(file_bytes)


def is_corrupted_pdf_text(text: str) -> bool:
    """Checks if extracted PDF text stream contains corrupted PUA glyphs or disoriented character tables."""
    if not text or not text.strip():
        return True
    
    pua_count = sum(1 for c in text if 0xE000 <= ord(c) <= 0xF8FF or 0xF0000 <= ord(c) <= 0x10FFFF)
    replacement_char_count = text.count("\ufffd") + text.count("")
    
    if pua_count >= 1 or replacement_char_count >= 1:
        return True
    
    # Check for legacy font garbage tokens
    if any(k in text for k in ['善', '周', '善', '啣', '呈', '吸', '咞', '吆']):
        return True
        
    return False


def parse_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF file.
    If extracted text stream is corrupted by legacy font CMap tables, renders pages to images and applies OCR.
    """
    extracted_text = ""
    
    # 1. Try PyMuPDF fitz text extraction
    if fitz:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages_text = []
            for page in doc:
                t = page.get_text()
                if t and t.strip():
                    pages_text.append(t.strip())
            extracted_text = "\n\n".join(pages_text)
        except Exception as e:
            print(f"fitz text extraction error: {e}")

    # 2. Fallback to pypdf text extraction if fitz is not available or empty
    if not extracted_text:
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            extracted_pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t and t.strip():
                    extracted_pages.append(t.strip())
            extracted_text = "\n\n".join(extracted_pages)
        except Exception as e:
            print(f"pypdf text extraction error: {e}")

    # 3. Check if extracted text stream is corrupted or contains legacy Al Kanz PUA glyphs
    if is_corrupted_pdf_text(extracted_text):
        print("Detected legacy custom font / PUA encoding in PDF. Running page rendering + OCR pipeline...")
        ocr_text = parse_pdf_via_ocr(file_bytes)
        if ocr_text and ocr_text.strip():
            return ocr_text.strip()

    return extracted_text.strip()


def parse_pdf_via_ocr(file_bytes: bytes) -> str:
    """Renders PDF pages to high-res PNG images and runs local offline OCR on each page."""
    try:
        from backend.ocr_engine import extract_text_from_image
        
        pages_text = []
        if fitz:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                res = extract_text_from_image(img_bytes)
                if res.get("success") and res.get("text"):
                    pages_text.append(res["text"].strip())
                    
        return "\n\n".join(pages_text)
    except Exception as e:
        print(f"PDF OCR rendering error: {e}")
        return ""


def parse_docx(file_bytes: bytes) -> str:
    """Extracts text paragraphs and table cell text from Word .docx file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text)
                
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
                    
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"Error parsing .docx file: {e}")
        return parse_txt(file_bytes)


def parse_txt(file_bytes: bytes) -> str:
    """Decodes text bytes trying common encodings (UTF-8, UTF-16, Windows-1256 for Arabic)."""
    encodings = ['utf-8', 'utf-16', 'windows-1256', 'latin-1']
    for enc in encodings:
        try:
            return file_bytes.decode(enc)
        except Exception:
            continue
    return file_bytes.decode('utf-8', errors='ignore')
