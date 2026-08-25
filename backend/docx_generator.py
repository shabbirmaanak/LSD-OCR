"""
Word (.docx) Document Generator for Lisan al Dawat (LSD) OCR System.
Configures native Right-To-Left (RTL) XML properties, Arabic font styling,
and paragraph structure so generated documents render cleanly in MS Word & LibreOffice.
"""

import io
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def create_element(name: str) -> OxmlElement:
    """Helper to create an OpenXML element."""
    return OxmlElement(name)


def set_element_rtl(element) -> None:
    """Sets w:rtl tag on an OpenXML element (rPr or pPr)."""
    rtl_elem = create_element('w:rtl')
    element.append(rtl_elem)


def set_paragraph_rtl(paragraph) -> None:
    """Enforces Right-To-Left paragraph properties and alignment in docx."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    
    # Enable bidi (bi-directional / RTL paragraph)
    bidi_elem = create_element('w:bidi')
    pPr.append(bidi_elem)


def set_run_arabic_font(run, font_name: str = "Amiri", font_size_pt: int = 14, is_bold: bool = False, is_header: bool = False):
    """Configures run font properties specifically for Arabic / Lisan al Dawat script."""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = is_bold
    
    # Set complex script (cs) font properties in XML
    rPr = run._r.get_or_add_rPr()
    
    # w:rFonts with w:cs and w:ascii/w:hAnsi
    rFonts = create_element('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)
    
    # Set RTL flag on the run
    rtl_elem = create_element('w:rtl')
    rPr.append(rtl_elem)
    
    if is_header:
        run.font.color.rgb = RGBColor(15, 76, 129)  # Bohra blue/teal accent header color


def generate_lsd_docx(
    pages_data: List[Dict[str, Any]],
    document_title: str = "Lisan al Dawat Extracted Document",
    include_images: bool = True,
    font_name: str = "Al-Kanz",
    font_size: int = 14
) -> bytes:
    """
    Generates a native Word (.docx) binary stream from OCR extracted pages.
    
    pages_data structure:
    [
        {
            "page_num": 1,
            "text": "Extracted Lisan al Dawat text...",
            "image_bytes": b"..." (optional)
        }, ...
    ]
    """
    doc = Document()
    
    # Set page margins to standard 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        # Enable Right-to-Left section flow if possible
        sectPr = section._sectPr
        bidi_sect = create_element('w:bidi')
        sectPr.append(bidi_sect)

    # Document Header / Title
    title_p = doc.add_paragraph()
    set_paragraph_rtl(title_p)
    title_run = title_p.add_run(document_title)
    set_run_arabic_font(title_run, font_name=font_name, font_size_pt=20, is_bold=True, is_header=True)
    
    # Bismillah / Header Banner line
    subtitle_p = doc.add_paragraph()
    set_paragraph_rtl(subtitle_p)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("بسم الله الرحمن الرحيم")
    set_run_arabic_font(sub_run, font_name=font_name, font_size_pt=18, is_bold=True, is_header=True)
    
    doc.add_paragraph()  # Blank spacing

    # Iterate over pages
    for i, page in enumerate(pages_data):
        page_num = page.get("page_num", i + 1)
        raw_text = page.get("text", "").strip()
        img_bytes = page.get("image_bytes")
        
        if i > 0:
            doc.add_page_break()
        
        # Page Section Header
        header_p = doc.add_paragraph()
        set_paragraph_rtl(header_p)
        h_run = header_p.add_run(f"الصفحة - {page_num}")
        set_run_arabic_font(h_run, font_name=font_name, font_size_pt=12, is_bold=True, is_header=False)
        header_p.paragraph_format.space_after = Pt(6)
        
        # Split page content into paragraphs
        lines = raw_text.split('\n')
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            
            p = doc.add_paragraph()
            set_paragraph_rtl(p)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_after = Pt(8)
            
            # Check if line looks like a title/heading
            is_heading = line_str.startswith("#") or len(line_str) < 40 and not line_str.endswith(".")
            
            clean_line = line_str.lstrip("#").strip()
            r = p.add_run(clean_line)
            set_run_arabic_font(
                r,
                font_name=font_name,
                font_size_pt=font_size + (2 if is_heading else 0),
                is_bold=is_heading
            )
            
        # Optionally attach scanned page image
        if include_images and img_bytes:
            try:
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_stream = io.BytesIO(img_bytes)
                img_p.add_run().add_picture(img_stream, width=Inches(4.5))
                
                caption_p = doc.add_paragraph()
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = caption_p.add_run(f"(صورة الصفحة الممسوحة ضوئياً {page_num})")
                set_run_arabic_font(cap_run, font_name=font_name, font_size_pt=10, is_bold=False)
            except Exception as e:
                print(f"Warning: Could not attach image for page {page_num}: {e}")

    # Save to binary buffer
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream.getvalue()
